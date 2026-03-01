"""
Production-grade resume text extraction with OCR fallback.
Extracts and normalizes resume text from PDF/DOCX files.
Uses LLM for intelligent data extraction with regex fallback.
"""

import io
import re
import logging
import pdfplumber
import docx
from fastapi import UploadFile
from typing import Tuple, Dict, Any

logger = logging.getLogger(__name__)

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text from PDF or DOCX file bytes.
    This is the main extraction function that works with raw bytes.
    """
    filename_lower = filename.lower()
    file_ext = filename_lower.split(".")[-1] if "." in filename_lower else ""

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf":
        extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short
        if OCR_AVAILABLE and (not extracted_text or len(extracted_text.strip()) < 100):
            extracted_text = _extract_from_pdf_ocr(content)
            if not extracted_text:
                raise ValueError("Failed to extract text from PDF (even with OCR)")
        elif not extracted_text:
            raise ValueError("Failed to extract text from PDF")

    # DOCX handling
    elif file_ext == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Please upload a PDF or DOCX file.")

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError(f"Resume text is too short or empty after extraction. Got {len(normalized_text)} chars.")
    
    return normalized_text


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from PDF or DOCX file.
    Falls back to OCR if PDF text extraction is empty.
    Normalizes and returns clean text.
    """
    content = await file.read()
    filename = (file.filename or "").lower()
    file_ext = filename.split(".")[-1] if "." in filename else ""
    content_type = (file.content_type or "").lower()

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf" or "pdf" in content_type:
        extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short
        if OCR_AVAILABLE and (not extracted_text or len(extracted_text.strip()) < 100):
            extracted_text = _extract_from_pdf_ocr(content)
            if not extracted_text:
                raise ValueError("Failed to extract text from PDF (even with OCR)")
        elif not extracted_text:
            raise ValueError("Failed to extract text from PDF")

    # DOCX handling
    elif file_ext == "docx" or "wordprocessingml" in content_type or "docx" in content_type:
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError("Resume text is too short or empty after extraction.")
    
    return normalized_text


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        return text
    except Exception:
        return ""


def _extract_from_pdf_ocr(content: bytes) -> str:
    """Extract text from PDF using OCR (pytesseract + pdf2image)."""
    if not OCR_AVAILABLE:
        return ""
    
    try:
        images = convert_from_bytes(content)
        text = ""
        for image in images:
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
        return text
    except Exception:
        return ""


def _normalize_text(text: str) -> str:
    """
    Normalize resume text:
    - Convert to lowercase
    - Remove extra whitespace
    - Preserve emails and URLs
    """
    # Preserve emails and URLs by replacing them temporarily
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    url_pattern = r'https?://[^\s]+'
    
    emails = re.findall(email_pattern, text)
    urls = re.findall(url_pattern, text)
    
    # Replace with placeholders (lowercase to survive lowercasing)
    text_with_placeholders = text
    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(email, f"__email_{i}__")
    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(url, f"__url_{i}__")
    
    # Normalize
    text_with_placeholders = text_with_placeholders.lower()
    text_with_placeholders = re.sub(r'\s+', ' ', text_with_placeholders)
    
    # Restore emails and URLs
    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(f"__email_{i}__", email)
    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(f"__url_{i}__", url)
    
    return text_with_placeholders.strip()


def extract_candidate_info(resume_text: str) -> dict:
    """
    Extract structured candidate information from resume text.
    Uses LLM as primary method with regex fallback.
    
    Returns: {
        name, email, phone, linkedin_url, github_url, portfolio_url,
        skills, experience_years, education, projects, certifications,
        summary, extraction_method
    }
    """
    # Try LLM extraction first
    try:
        from app.services.llm_extractor import LLMResumeExtractor
        
        llm_extractor = LLMResumeExtractor()
        extracted_data = llm_extractor.extract_resume_data(resume_text)
        
        # Add extraction method for debugging
        extracted_data["extraction_method"] = "llm" if extracted_data.get("name") != "Unknown" else "regex"
        
        logger.info(f"Resume extraction completed using: {extracted_data['extraction_method']}")
        return extracted_data
        
    except Exception as e:
        logger.warning(f"LLM extraction failed, using regex fallback: {e}")
        return _extract_candidate_info_regex(resume_text)


def _extract_candidate_info_regex(resume_text: str) -> dict:
    """
    Fallback regex-based extraction of candidate information.
    """
    info = {
        "name": None,
        "email": None,
        "phone": None,
        "linkedin_url": None,
        "github_url": None,
        "skills": [],
        "experience_years": 0,
        "experience_months": 0,
        "education": [],
        "certifications": [],
        "summary": None,
        "extraction_method": "regex"
    }
    
    text = resume_text.lower()
    original_text = resume_text
    
    # Email extraction - improved pattern to handle more formats
    email_patterns = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # Standard email
        r'\b[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # With spaces around @
        r'\b[A-Za-z0-9._%+-]+\s*\[\s*at\s*\]\s*[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # [at] format
    ]
    
    for pattern in email_patterns:
        email_match = re.search(pattern, original_text, re.IGNORECASE)
        if email_match:
            email = email_match.group().replace(' ', '').replace('[at]', '@')
            info["email"] = email
            break
    
    # Name extraction - improved to handle various formats
    # Strategy 1: First extract text before special characters/contact info
    name_text = original_text
    
    # Common separators in resumes that come after name
    separators = [
        r'\s*\+\s*',  # + symbol
        r'\s*#\s*',   # # symbol
        r'\s*\|\s*',  # | symbol
        r'\s*•\s*',   # bullet
        r'\s*@\s*',   # @ for email
        r'linkedin',
        r'github',
        r'\d{10}',    # phone number
        r'\+\d{1,3}[-\s]?\d',  # international phone
    ]
    
    for sep in separators:
        match = re.search(sep, name_text, re.IGNORECASE)
        if match:
            name_text = name_text[:match.start()].strip()
            break
    
    # Clean up the extracted name
    name_text = name_text.strip()
    
    # Check if it's a valid name (2-5 words, mostly letters, no special chars)
    if name_text:
        words = name_text.split()
        if 1 <= len(words) <= 5:
            # Verify it looks like a name
            name_candidate = ' '.join(words[:5])
            # Remove any remaining special characters
            name_candidate = re.sub(r'[^\w\s]', '', name_candidate).strip()
            
            if name_candidate and len(name_candidate) >= 3:
                # Check letter ratio
                letter_count = sum(1 for c in name_candidate if c.isalpha() or c.isspace())
                if letter_count / len(name_candidate) > 0.8:
                    # Title case the name
                    info["name"] = name_candidate.title()
    
    # Strategy 2: If name not found, try newline-based extraction
    if not info["name"]:
        lines = original_text.split('\n')
        for line in lines[:5]:
            line_clean = line.strip()
            if (line_clean and 
                3 <= len(line_clean) <= 50 and 
                1 <= len(line_clean.split()) <= 5 and
                not any(keyword in line_clean.lower() for keyword in ['email', 'phone', 'linkedin', 'github', 'http', ':', '@', '#', '+'])):
                
                letter_ratio = sum(1 for c in line_clean if c.isalpha() or c.isspace()) / len(line_clean)
                if letter_ratio > 0.7:
                    info["name"] = line_clean.title()
                    break
    
    # Phone extraction (common patterns)
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\(\d{3}\)\s\d{3}-\d{4}',
        r'\d{3}-\d{3}-\d{4}'
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, resume_text)
        if phone_match:
            info["phone"] = phone_match.group()
            break
    
    # LinkedIn URL extraction
    linkedin_match = re.search(r'linkedin\.com/in/[^\s]+', text)
    if linkedin_match:
        info["linkedin_url"] = "https://" + linkedin_match.group()
    
    # GitHub URL extraction
    github_match = re.search(r'github\.com/[^\s]+', text)
    if github_match:
        info["github_url"] = "https://" + github_match.group()
    
    # Years of experience extraction
    # First try explicit "X years" patterns
    years_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)[:\s]+(\d+)\+?\s*(?:years?|yrs?)',
        r'total\s+(?:experience|exp)[:\s]+(\d+)',
    ]
    
    experience_years = 0.0
    total_experience_months = 0
    
    for pattern in years_patterns:
        matches = re.findall(pattern, text)
        if matches:
            years_found = [int(m) for m in matches if m.strip()]
            if years_found:
                experience_years = max(years_found)
                total_experience_months = int(experience_years * 12)
                break
    
    # If no explicit years found, try to calculate from date ranges
    if experience_years == 0:
        from datetime import datetime
        
        # Multiple date range patterns
        date_range_patterns = [
            r'([a-z]{3,9})\s*[\'"]?(\d{4})\s*[–\-–—to]+\s*([a-z]{3,9}|present|current|now)\s*[\'"]?(\d{4})?',
            r'(\d{2})/(\d{4})\s*[–\-–—to]+\s*(\d{2}|present|current)/(\d{4})?',
        ]
        
        months = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'sept': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }
        
        current_dt = datetime.now()
        total_months = 0
        
        for pattern in date_range_patterns:
            date_matches = re.findall(pattern, text, re.IGNORECASE)
            
            for match in date_matches:
                try:
                    start_month_str = match[0].lower()[:3]
                    start_year = int(match[1])
                    end_month_str = match[2].lower()[:3] if match[2] else ''
                    end_year = int(match[3]) if match[3] and match[3].isdigit() else None
                    
                    # Get start month
                    start_month = months.get(start_month_str, 1)
                    
                    # Handle "present" or current date
                    if end_month_str in ['pre', 'cur', 'now'] or not end_month_str:
                        end_year = current_dt.year
                        end_month = current_dt.month
                    else:
                        end_month = months.get(end_month_str, 12)
                        if not end_year:
                            end_year = start_year
                    
                    # Calculate months
                    if start_year and end_year:
                        calc_months = (end_year - start_year) * 12 + (end_month - start_month)
                        if 0 < calc_months < 240:  # Max 20 years
                            total_months += calc_months
                except (ValueError, TypeError, IndexError):
                    continue
        
        if total_months > 0:
            total_experience_months = total_months
            experience_years = round(total_months / 12, 2)
    
    info["experience_years"] = experience_years
    info["experience_months"] = total_experience_months
    
    # Education degree extraction - extract EXACT degree names from resume
    # Map patterns to actual degree names found in text
    education_patterns = [
        # Full degree names with field
        (r'\b(master\s+of\s+computer\s+application|mca)\b', 'MCA (Master of Computer Application)'),
        (r'\b(bachelor\s+of\s+computer\s+application|bca)\b', 'BCA (Bachelor of Computer Application)'),
        (r'\b(m\.?tech|master\s+of\s+technology)\b', 'M.Tech'),
        (r'\b(b\.?tech|bachelor\s+of\s+technology)\b', 'B.Tech'),
        (r'\b(m\.?sc|master\s+of\s+science)\b', 'M.Sc'),
        (r'\b(b\.?sc|bachelor\s+of\s+science)\b', 'B.Sc'),
        (r'\b(mba|master\s+of\s+business\s+administration)\b', 'MBA'),
        (r'\b(b\.?e\.?|bachelor\s+of\s+engineering)\b', 'B.E.'),
        (r'\b(m\.?e\.?|master\s+of\s+engineering)\b', 'M.E.'),
        (r'\b(ph\.?d\.?|doctorate|doctoral)\b', 'Ph.D.'),
        (r'\b(b\.?a\.?|bachelor\s+of\s+arts)\b', 'B.A.'),
        (r'\b(m\.?a\.?|master\s+of\s+arts)\b', 'M.A.'),
        (r'\b(b\.?com|bachelor\s+of\s+commerce)\b', 'B.Com'),
        (r'\b(m\.?com|master\s+of\s+commerce)\b', 'M.Com'),
        (r'\bhsc\b', 'HSC (Higher Secondary)'),
        (r'\bssc\b', 'SSC (Secondary School)'),
        (r'\b(diploma|associate)\b', 'Diploma'),
    ]
    
    education_found = []
    for pattern, degree_name in education_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            if degree_name not in education_found:
                education_found.append(degree_name)
    
    info["education"] = education_found
    
    # Skills extraction
    skill_keywords = [
        'python', 'javascript', 'java', 'c++', 'c#', 'rust', 'go', 'kotlin', 'swift',
        'react', 'angular', 'vue', 'node', 'express', 'django', 'flask', 'fastapi',
        'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins',
        'html', 'css', 'sql', 'git', 'linux', 'windows',
        'machine learning', 'nlp', 'computer vision', 'deep learning',
        'rest api', 'graphql', 'websocket', 'microservices', 'devops',
        'agile', 'scrum', 'jira', 'confluence', 'slack'
    ]
    for skill in skill_keywords:
        if skill in text:
            info["skills"].append(skill.title())
    info["skills"] = list(set(info["skills"]))
    
    # Clean the name if extracted
    if info["name"]:
        info["name"] = _clean_extracted_name(info["name"])
    
    return info


def _clean_extracted_name(name: str) -> str:
    """Clean extracted name - remove address, symbols, etc."""
    if not name or name == "Unknown":
        return "Unknown"
    
    # Stop at common separators
    separators = ['+', '#', '|', '@', '(', ')', ',', ':', ';']
    for sep in separators:
        if sep in name:
            name = name.split(sep)[0]
    
    # Remove numbers (likely address or phone)
    name = re.sub(r'\d+', '', name)
    
    # Remove extra whitespace and clean
    name = re.sub(r'\s+', ' ', name).strip()
    
    # Remove common non-name words
    non_name_words = ['l-', 'apt', 'street', 'road', 'nagar', 'vadodara', 'india', 'address', 'taif']
    name_lower = name.lower()
    for word in non_name_words:
        if word in name_lower:
            idx = name_lower.find(word)
            name = name[:idx].strip()
            name_lower = name.lower()
    
    # Ensure it looks like a name (mostly letters)
    if name:
        letter_count = sum(1 for c in name if c.isalpha() or c.isspace())
        if len(name) > 0 and letter_count / len(name) < 0.8:
            return "Unknown"
    
    # Title case and limit to reasonable length
    name = name.strip().title()
    words = name.split()
    if len(words) > 5:
        name = ' '.join(words[:5])
    
    return name if len(name) >= 2 else "Unknown"
