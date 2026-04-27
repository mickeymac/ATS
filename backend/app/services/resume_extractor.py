"""
Production-grade resume text extraction with OCR fallback.
Extracts and normalizes resume text from PDF/DOCX files.
Uses LLM for intelligent data extraction with regex fallback.
"""

import io
import re
import logging
import easyocr
import numpy as np
from PIL import Image
import shutil
import subprocess
import pdfplumber
import docx
from fastapi import UploadFile
from typing import Tuple, Dict, Any

logger = logging.getLogger(__name__)

# Check for pymupdf
try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Check for EasyOCR (pure Python, no external dependencies)
EASYOCR_AVAILABLE = False
_easyocr_reader = None
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    logger.info("EasyOCR available for scanned PDF extraction")
except ImportError:
    logger.warning("EasyOCR not available")

# Check for pytesseract OCR  
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    # Actually check if Tesseract binary is installed
    tesseract_path = shutil.which("tesseract")
    if tesseract_path:
        OCR_AVAILABLE = True
        logger.info(f"Tesseract OCR found at: {tesseract_path}")
    else:
        OCR_AVAILABLE = False
        logger.warning("pytesseract installed but Tesseract binary not found in PATH")
except ImportError:
    OCR_AVAILABLE = False

# Check for pymupdf OCR capability
PYMUPDF_OCR_AVAILABLE = False
if PYMUPDF_AVAILABLE:
    try:
        # pymupdf can do OCR if Tesseract is installed
        result = subprocess.run(["tesseract", "--version"], capture_output=True, timeout=5)
        if result.returncode == 0:
            PYMUPDF_OCR_AVAILABLE = True
            logger.info("pymupdf OCR capability available")
    except Exception:
        pass


def _get_easyocr_reader():
    """Lazy-load EasyOCR reader to avoid slow startup."""
    global _easyocr_reader
    if _easyocr_reader is None and EASYOCR_AVAILABLE:
        import easyocr
        _easyocr_reader = easyocr.Reader(['en'], gpu=False)  # Use CPU for compatibility
        logger.info("EasyOCR reader initialized")

        print("_get_easyocr_reader: EasyOCR reader initialized")
    return _easyocr_reader


 #Image OCR

reader = easyocr.Reader(['en'])

def _extract_from_image_easyocr(content: bytes):
    image = Image.open(io.BytesIO(content))
    image_np = np.array(image)
    results = reader.readtext(image_np, detail=0)
    print("_extract_from_image_easyocr: Initialized")
    return "\n".join(results)
    
def _extract_from_image_tesseract(content: bytes):
    """
    OCR fallback using pytesseract for image files
    """
    image = Image.open(io.BytesIO(content))
    text = pytesseract.image_to_string(image)
    print("_extract_from_image_tesseract: OCR completed")
    return text

# def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text from PDF or DOCX file bytes.
    This is the main extraction function that works with raw bytes.
    """
    filename_lower = filename.lower()
    file_ext = filename_lower.split(".")[-1] if "." in filename_lower else ""

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf":
        # Try pymupdf first (faster and more reliable)
        if PYMUPDF_AVAILABLE:
            extracted_text = _extract_from_pdf_pymupdf(content)
        
        # Fallback to pdfplumber if pymupdf failed
        if not extracted_text or len(extracted_text.strip()) < 50:
            extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short (likely scanned PDF)
        if not extracted_text or len(extracted_text.strip()) < 100:
            logger.info("PDF appears to be scanned, attempting OCR extraction...")
            
            # Try EasyOCR first (no external dependencies)
            if EASYOCR_AVAILABLE:
                logger.info("Using EasyOCR for scanned PDF...")
                extracted_text = _extract_from_pdf_easyocr(content)
            
            # Try pymupdf OCR
            if (not extracted_text or len(extracted_text.strip()) < 50) and PYMUPDF_OCR_AVAILABLE:
                logger.info("Trying pymupdf OCR...")
                extracted_text = _extract_from_pdf_pymupdf_ocr(content)
            
            # Try pytesseract OCR
            if (not extracted_text or len(extracted_text.strip()) < 50) and OCR_AVAILABLE:
                logger.info("Trying pytesseract OCR...")
                extracted_text = _extract_from_pdf_ocr(content)
            
            if not extracted_text or len(extracted_text.strip()) < 50:
                raise ValueError(
                    "This PDF appears to be a scanned image and OCR extraction failed. "
                    "Please upload a text-based PDF or DOCX file instead."
                )
            

    # DOCX handling
    elif file_ext == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Please upload a PDF or DOCX file.")

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError(f"Resume text is too short or empty after extraction. Got {len(normalized_text)} chars.")
    
    return normalized_text

def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text from PDF, DOCX, or Image file bytes.
    """
    print("Starting text extraction from bytes...")
    filename_lower = filename.lower()
    file_ext = filename_lower.split(".")[-1] if "." in filename_lower else ""

    extracted_text = ""

    # ---------------- PDF handling ----------------
    if file_ext == "pdf":
        if PYMUPDF_AVAILABLE:
            extracted_text = _extract_from_pdf_pymupdf(content)

        if not extracted_text or len(extracted_text.strip()) < 50:
            extracted_text = _extract_from_pdf(content)

        if not extracted_text or len(extracted_text.strip()) < 100:
            logger.info("PDF appears to be scanned, attempting OCR extraction...")

            if EASYOCR_AVAILABLE:
                logger.info("Using EasyOCR for scanned PDF...")
                extracted_text = _extract_from_pdf_easyocr(content)

            if (not extracted_text or len(extracted_text.strip()) < 50) and PYMUPDF_OCR_AVAILABLE:
                logger.info("Trying pymupdf OCR...")
                extracted_text = _extract_from_pdf_pymupdf_ocr(content)

            if (not extracted_text or len(extracted_text.strip()) < 50) and OCR_AVAILABLE:
                logger.info("Trying pytesseract OCR...")
                extracted_text = _extract_from_pdf_ocr(content)

            if not extracted_text or len(extracted_text.strip()) < 50:
                raise ValueError(
                    "This PDF appears to be a scanned image and OCR extraction failed. "
                    "Please upload a text-based PDF or DOCX file instead."
                )

    # ---------------- DOCX handling ----------------
    elif file_ext == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    # ---------------- IMAGE handling ----------------
    elif file_ext in ["png", "jpg", "jpeg", "bmp", "tiff", "webp"]:
        logger.info("Processing image file for OCR extraction...")
       

        # Try EasyOCR first
        if EASYOCR_AVAILABLE:
            try:
                extracted_text = _extract_from_image_easyocr(content)
            except Exception:
                extracted_text = ""

        # Fallback to pytesseract
        if (not extracted_text or len(extracted_text.strip()) < 20) and OCR_AVAILABLE:
            try:
                extracted_text = _extract_from_image_tesseract(content)
            except Exception:
                extracted_text = ""

        if not extracted_text or len(extracted_text.strip()) < 20:
            raise ValueError(
                "Text extraction from image failed. Please upload a clearer image or a PDF/DOCX file."
            )

    # ---------------- Unsupported ----------------
    else:
        raise ValueError(
            f"Unsupported file format: {file_ext}. Please upload PDF, DOCX, or Image files."
        )

    # return extracted_text

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError(f"Resume text is too short or empty after extraction. Got {len(normalized_text)} chars.")
    
    return normalized_text

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from PDF or DOCX file.
    Falls back to OCR if PDF text extraction is empty.
    Normalizes and returns clean text.
    """
    print("Starting text extraction from file...")
    content = await file.read()
    filename = (file.filename or "").lower()
    file_ext = filename.split(".")[-1] if "." in filename else ""
    content_type = (file.content_type or "").lower()

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf" or "pdf" in content_type:
        # Try pymupdf first (faster and more reliable)
        if PYMUPDF_AVAILABLE:
            extracted_text = _extract_from_pdf_pymupdf(content)
        
        # Fallback to pdfplumber if pymupdf failed
        if not extracted_text or len(extracted_text.strip()) < 50:
            extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short (likely scanned PDF)
        if not extracted_text or len(extracted_text.strip()) < 100:
            logger.info("PDF appears to be scanned, attempting OCR extraction...")
            
            # Try EasyOCR first (no external dependencies)
            if EASYOCR_AVAILABLE:
                logger.info("Using EasyOCR for scanned PDF...")
                extracted_text = _extract_from_pdf_easyocr(content)
            
            # Try pymupdf OCR
            if (not extracted_text or len(extracted_text.strip()) < 50) and PYMUPDF_OCR_AVAILABLE:
                logger.info("Trying pymupdf OCR...")
                extracted_text = _extract_from_pdf_pymupdf_ocr(content)
            
            # Try pytesseract OCR
            if (not extracted_text or len(extracted_text.strip()) < 50) and OCR_AVAILABLE:
                logger.info("Trying pytesseract OCR...")
                extracted_text = _extract_from_pdf_ocr(content)
            
            if not extracted_text or len(extracted_text.strip()) < 50:
                raise ValueError(
                    "This PDF appears to be a scanned image and OCR extraction failed. "
                    "Please upload a text-based PDF or DOCX file instead."
                )

    # DOCX handling
    elif file_ext == "docx" or "wordprocessingml" in content_type or "docx" in content_type:
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    elif file_ext in ["png", "jpg", "jpeg", "bmp", "tiff", "webp"] or "image" in content_type:
        logger.info("Processing image file for OCR extraction...")

    # Try EasyOCR first
    if EASYOCR_AVAILABLE:
        try:
            extracted_text = _extract_from_image_easyocr(content)
        except Exception:
            extracted_text = ""

    # Fallback to pytesseract
    if (not extracted_text or len(extracted_text.strip()) < 20) and OCR_AVAILABLE:
        try:
            extracted_text = _extract_from_image_tesseract(content)
        except Exception:
            extracted_text = ""

    if not extracted_text or len(extracted_text.strip()) < 20:
        raise ValueError(
            "Text extraction from image failed. Please upload a clearer image or a PDF/DOCX file."
        )
    
    else:
        raise ValueError("Unsupported file format. Please upload PDF, DOCX, or Image files.")
    
    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError("Resume text is too short or empty after extraction.")
    
    return normalized_text


def extract_profile_picture_from_pdf(content: bytes) -> bytes:
    """
    Extracts the most likely profile picture from the first page of a PDF.
    Returns the image bytes or None.
    """
    if not PYMUPDF_AVAILABLE:
        return None

    try:
        doc = fitz.open(stream=content, filetype="pdf")
        if len(doc) == 0:
            return None
        
        # Only look at the first page for profile pictures
        page = doc[0]
        image_list = page.get_images(full=True)
        
        best_image_bytes = None
        best_image_score = -1

        for img in image_list:
            xref = img[0]
            base_image = doc.extract_image(xref)
            if not base_image:
                continue
            
            image_bytes = base_image["image"]
            width = base_image["width"]
            height = base_image["height"]
            
            # Filter out tiny icons and massive background graphics
            if width < 80 or height < 80 or width > 800 or height > 800:
                continue
                
            # Profile pictures are usually somewhat square / portrait
            aspect_ratio = width / height
            if aspect_ratio < 0.5 or aspect_ratio > 1.5:
                continue
                
            # Score based on size (prefer decently sized images)
            score = width * height
            if score > best_image_score:
                best_image_score = score
                best_image_bytes = image_bytes
                
        doc.close()
        return best_image_bytes
    except Exception as e:
        logger.warning(f"Profile picture extraction failed: {e}")
        return None

def _extract_from_pdf_pymupdf(content: bytes) -> str:
    """Extract text from PDF using pymupdf (fitz) - faster than pdfplumber."""
    print("_extract_from_pdf_pymupdf: Starting extraction with pymupdf...")
    if not PYMUPDF_AVAILABLE:
        return ""
    try:
        text = ""
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            text += page_text + "\n"
        doc.close()
        return text
    except Exception as e:
        logger.warning(f"pymupdf extraction failed: {e}")
        return ""


def _extract_from_pdf_pymupdf_ocr(content: bytes) -> str:
    """Extract text from scanned PDF using pymupdf OCR."""
    print("_extract_from_pdf_pymupdf_ocr: Starting OCR extraction...")
    if not PYMUPDF_AVAILABLE:
        return ""
    try:
        text = ""
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            # Use pymupdf's OCR capability
            page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            if not page_text or len(page_text.strip()) < 20:
                # Try OCR if direct text extraction failed
                try:
                    page_text = page.get_text(textpage=None, option=fitz.TEXT_OUTPUT_TEXT)
                except Exception:
                    pass
            text += page_text + "\n"
        doc.close()
        return text
    except Exception as e:
        logger.warning(f"pymupdf OCR extraction failed: {e}")
        return ""


def _extract_from_pdf_easyocr(content: bytes) -> str:
    """Extract text from scanned PDF using EasyOCR (pure Python, no external binaries)."""
    print("_extract_from_pdf_easyocr: Starting EasyOCR extraction...")
    if not EASYOCR_AVAILABLE or not PYMUPDF_AVAILABLE:
        return ""
    
    try:
        import numpy as np
        from PIL import Image
        
        reader = _get_easyocr_reader()
        if reader is None:
            return ""
        
        text = ""
        doc = fitz.open(stream=content, filetype="pdf")
        
        for page_num, page in enumerate(doc):
            logger.info(f"Processing page {page_num + 1}/{len(doc)} with EasyOCR...")
            
            # Convert page to image (high DPI for better OCR)
            mat = fitz.Matrix(2, 2)  # 2x zoom for better resolution
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to numpy array for EasyOCR
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            img_array = np.array(img)
            
            # Run OCR
            results = reader.readtext(img_array, detail=0, paragraph=True)
            page_text = "\n".join(results)
            text += page_text + "\n\n"
        
        doc.close()
        logger.info(f"EasyOCR extracted {len(text)} characters")
        return text
        
    except Exception as e:
        logger.warning(f"EasyOCR extraction failed: {e}")
        return ""


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    print("_extract_from_pdf: Starting extraction with pdfplumber...")
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        return text
    except Exception:
        return ""


def _extract_from_pdf_ocr(content: bytes) -> str:
    """Extract text from PDF using OCR (pytesseract + pdf2image)."""
    print("_extract_from_pdf_ocr: Starting OCR extraction with pytesseract...")
    if not OCR_AVAILABLE:
        return ""
    
    try:
        images = convert_from_bytes(content)
        text = ""
        for image in images:
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
        return text
    except Exception:
        return ""


"""def _normalize_text(text: str) -> str:
    
    #Normalize resume text:
    - #Convert to lowercase
    - #Remove extra whitespace
    - #Preserve emails and URLs
    
    #Phone Number Normalization - fix common OCR errors
    text = _fix_ocr_numbers(phone)

    print("_normalize_text: Starting text normalization...")
    # Preserve emails and URLs by replacing them temporarily
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    url_pattern = r'https?://[^\s]+'
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\(\d{3}\)\s?\d{3}-\d{4}',
        r'\d{3}-\d{3}-\d{4}',
    ]
    
    emails = re.findall(email_pattern, text)
    urls = re.findall(url_pattern, text)
    for pattern in phone_patterns:
        phone = re.findall(pattern, text)
        if phone:
            info["phone"] = phone[0]
    
    # Replace with placeholders (lowercase to survive lowercasing)
    text_with_placeholders = text
    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(email, f"__email_{i}__")
    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(url, f"__url_{i}__")
    for i, phone in enumerate(phone):
        text_with_placeholders = text_with_placeholders.replace(phone, f"__phone_{i}__")
    
    # Normalize
    text_with_placeholders = text_with_placeholders.lower()
    text_with_placeholders = re.sub(r'\s+', ' ', text_with_placeholders)
    
    # Restore emails and URLs
    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(f"__email_{i}__", email)
    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(f"__url_{i}__", url)
    for i, phone in enumerate(phone):
        text_with_placeholders = text_with_placeholders.replace(f"__phone_{i}__", phone)



    
    return text_with_placeholders.strip()"""

def _normalize_text(text: str) -> str:
    """
    Normalize resume text:
    - Fix OCR number errors
    - Convert to lowercase
    - Remove extra whitespace
    - Preserve emails, URLs, and phone numbers
    """

    print("_normalize_text: Starting text normalization...")
    print("Original text length:", len(text),text)

    if not text:
        return ""

    # Fix OCR errors in numbers

    # Patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    url_pattern = r'https?://[^\s]+'
    # phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?[\dOIlSB]{3}\)?[-.\s]?[\dOIlSB]{3}[-.\s]?[\dOIlSB]{4}'
    phone_pattern = r'\b[\dA-Za-z\]\|]{2,4}[\s\-][\dA-Za-z\]\|]{2,4}-[\dA-Za-z\]\|]{3,4}\b'

    # Extract items
    emails = re.findall(email_pattern, text)
    urls = re.findall(url_pattern, text)
    testphones = re.findall(phone_pattern, text)
    print("Found phones (raw):", testphones)
    phones = []

    for match in re.findall(phone_pattern, text):
        fixed = _fix_ocr_numbers(match)
        phones.append(fixed)
        print(f"Found phone: {match} -> Fixed: {fixed} -> Phones list: {phones}")

    print("Phones:", phones)

    # Replace with placeholders
    text_with_placeholders = text

    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(email, f"__EMAIL_{i}__")

    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(url, f"__URL_{i}__")

    for i, phone in enumerate(phones):
        text_with_placeholders = text_with_placeholders.replace(phone, f"__PHONE_{i}__")

    # Normalize text
    text_with_placeholders = text_with_placeholders.lower()
    text_with_placeholders = re.sub(r'\s+', ' ', text_with_placeholders)

    # Restore placeholders
    for i, email in enumerate(emails):
        text_with_placeholders = text_with_placeholders.replace(f"__email_{i}__", email)

    for i, url in enumerate(urls):
        text_with_placeholders = text_with_placeholders.replace(f"__url_{i}__", url)

    for i, phone in enumerate(phones):
        text_with_placeholders = text_with_placeholders.replace(f"__phone_{i}__", phone)

    return text_with_placeholders.strip()

#  def extract_candidate_info(resume_text: str) -> dict:
    """
    Extract structured candidate information from resume text.
    Uses LLM as primary method with regex fallback.
    
    Returns: {
        name, email, phone, linkedin_url, github_url, portfolio_url,
        skills, experience_years, education, projects, certifications,
        summary, extraction_method
    }
    
    # Try LLM extraction first
    try:
        from app.services.llm_extractor import LLMResumeExtractor
        print(resume_text)
        
        llm_extractor = LLMResumeExtractor()
        extracted_data = llm_extractor.extract_resume_data(resume_text)
        
        # Add extraction method for debugging
        extracted_data["extraction_method"] = "llm" if extracted_data.get("name") != "Unknown" else "regex"
        
        logger.info(f"Resume extraction completed using: {extracted_data['extraction_method']}")
        return extracted_data
        
    except Exception as e:
        logger.warning(f"LLM extraction failed, using regex fallback: {e}")
        return _extract_candidate_info_regex(resume_text)"""

#OCR: Fix Numbers
def _fix_ocr_numbers(text: str) -> str:
    print("In _fix_ocr_numbers")
    corrections = {
        "j": "3",
        "J": "3",
        "]": "1",
        "l": "1",
        "I": "1",
        "O": "0",
        "S": "5",
        "B": "8"
    }

    for wrong, correct in corrections.items():
        text = text.replace(wrong, correct)

    print("In _fix_ocr_numbers, original text:", text)
    return text

def _extract_candidate_info_regex(resume_text: str) -> dict:
    """
    Fallback regex-based extraction of candidate information.
    """
    print("_extract_candidate_info_regex: Starting regex extraction...")
    info = {
        "name": None,
        "email": None,
        "phone": None,
        "linkedin_url": None,
        "github_url": None,
        "skills": [],
        "experience_years": 0,
        "experience_months": 0,
        "education": [],
        "certifications": [],
        "summary": None,
        "extraction_method": "regex"
    }
    
    text = resume_text.lower()
    original_text = resume_text
    
    # Email extraction - improved pattern to handle more formats
    email_patterns = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # Standard email
        r'\b[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # With spaces around @
        r'\b[A-Za-z0-9._%+-]+\s*\[\s*at\s*\]\s*[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',  # [at] format
    ]
    
    for pattern in email_patterns:
        email_match = re.search(pattern, original_text, re.IGNORECASE)
        if email_match:
            email = email_match.group().replace(' ', '').replace('[at]', '@')
            info["email"] = email
            break
    
    # Name extraction - improved to handle various formats
    # Strategy 1: First extract text before special characters/contact info
    name_text = original_text
    
    # Common separators in resumes that come after name
    separators = [
        r'\s*\+\s*',  # + symbol
        r'\s*#\s*',   # # symbol
        r'\s*\|\s*',  # | symbol
        r'\s*•\s*',   # bullet
        r'\s*@\s*',   # @ for email
        r'linkedin',
        r'github',
        r'\d{10}',    # phone number
        r'\+\d{1,3}[-\s]?\d',  # international phone
    ]
    
    for sep in separators:
        match = re.search(sep, name_text, re.IGNORECASE)
        if match:
            name_text = name_text[:match.start()].strip()
            break
    
    # Clean up the extracted name
    name_text = name_text.strip()
    
    # Check if it's a valid name (2-5 words, mostly letters, no special chars)
    if name_text:
        words = name_text.split()
        if 1 <= len(words) <= 5:
            # Verify it looks like a name
            name_candidate = ' '.join(words[:5])
            # Remove any remaining special characters
            name_candidate = re.sub(r'[^\w\s]', '', name_candidate).strip()
            
            if name_candidate and len(name_candidate) >= 3:
                # Check letter ratio
                letter_count = sum(1 for c in name_candidate if c.isalpha() or c.isspace())
                if letter_count / len(name_candidate) > 0.8:
                    # Title case the name
                    info["name"] = name_candidate.title()
    
    # Strategy 2: If name not found, try newline-based extraction
    if not info["name"]:
        lines = original_text.split('\n')
        for line in lines[:5]:
            line_clean = line.strip()
            if (line_clean and 
                3 <= len(line_clean) <= 50 and 
                1 <= len(line_clean.split()) <= 5 and
                not any(keyword in line_clean.lower() for keyword in ['email', 'phone', 'linkedin', 'github', 'http', ':', '@', '#', '+'])):
                
                letter_ratio = sum(1 for c in line_clean if c.isalpha() or c.isspace()) / len(line_clean)
                if letter_ratio > 0.7:
                    info["name"] = line_clean.title()
                    break
    

    

# Fix OCR errors in numbers before phone extraction
    text = _fix_ocr_numbers(original_text)

# Phone extraction patterns
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\(\d{3}\)\s?\d{3}-\d{4}',
        r'\d{3}-\d{3}-\d{4}',
    ]

    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            info["phone"] = matches[0]
            break

    # Phone extraction (common patterns)
    """phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\(\d{3}\)\s\d{3}-\d{4}',
        r'\d{3}-\d{3}-\d{4}',
        r'(\+?\d{1,3}[- ]?)?\d{10}'
    ]"""


    # for pattern in phone_patterns:
    # phone_match = re.search(pattern, text)
    # if phone_match:
    #     info["phone"] = phone_match.group()
    #     break
    
    # LinkedIn URL extraction
    linkedin_match = re.search(r'linkedin\.com/in/[^\s]+', text)
    if linkedin_match:
        info["linkedin_url"] = "https://" + linkedin_match.group()
    
    # GitHub URL extraction
    github_match = re.search(r'github\.com/[^\s]+', text)
    if github_match:
        info["github_url"] = "https://" + github_match.group()
    
    # Years of experience extraction
    # First try explicit "X years" patterns
    years_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)[:\s]+(\d+)\+?\s*(?:years?|yrs?)',
        r'total\s+(?:experience|exp)[:\s]+(\d+)',
    ]
    
    experience_years = 0.0
    total_experience_months = 0
    
    for pattern in years_patterns:
        matches = re.findall(pattern, text)
        if matches:
            years_found = [int(m) for m in matches if m.strip()]
            if years_found:
                experience_years = max(years_found)
                total_experience_months = int(experience_years * 12)
                break
    
    # If no explicit years found, try to calculate from date ranges
    if experience_years == 0:
        from datetime import datetime
        
        # Multiple date range patterns
        date_range_patterns = [
            r'([a-z]{3,9})\s*[\'"]?(\d{4})\s*[–\-–—to]+\s*([a-z]{3,9}|present|current|now)\s*[\'"]?(\d{4})?',
            r'(\d{2})/(\d{4})\s*[–\-–—to]+\s*(\d{2}|present|current)/(\d{4})?',
        ]
        
        months = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'sept': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }
        
        current_dt = datetime.now()
        total_months = 0
        
        for pattern in date_range_patterns:
            date_matches = re.findall(pattern, text, re.IGNORECASE)
            
            for match in date_matches:
                try:
                    start_month_str = match[0].lower()[:3]
                    start_year = int(match[1])
                    end_month_str = match[2].lower()[:3] if match[2] else ''
                    end_year = int(match[3]) if match[3] and match[3].isdigit() else None
                    
                    # Get start month
                    start_month = months.get(start_month_str, 1)
                    
                    # Handle "present" or current date
                    if end_month_str in ['pre', 'cur', 'now'] or not end_month_str:
                        end_year = current_dt.year
                        end_month = current_dt.month
                    else:
                        end_month = months.get(end_month_str, 12)
                        if not end_year:
                            end_year = start_year
                    
                    # Calculate months
                    if start_year and end_year:
                        calc_months = (end_year - start_year) * 12 + (end_month - start_month)
                        if 0 < calc_months < 240:  # Max 20 years
                            total_months += calc_months
                except (ValueError, TypeError, IndexError):
                    continue
        
        if total_months > 0:
            total_experience_months = total_months
            experience_years = round(total_months / 12, 2)
    
    info["experience_years"] = experience_years
    info["experience_months"] = total_experience_months
    
    # Education degree extraction - extract EXACT degree names from resume
    # Map patterns to actual degree names found in text
    education_patterns = [
        # Full degree names with field
        (r'\b(master\s+of\s+computer\s+application|mca)\b', 'MCA (Master of Computer Application)'),
        (r'\b(bachelor\s+of\s+computer\s+application|bca)\b', 'BCA (Bachelor of Computer Application)'),
        (r'\b(m\.?tech|master\s+of\s+technology)\b', 'M.Tech'),
        (r'\b(b\.?tech|bachelor\s+of\s+technology)\b', 'B.Tech'),
        (r'\b(m\.?sc|master\s+of\s+science)\b', 'M.Sc'),
        (r'\b(b\.?sc|bachelor\s+of\s+science)\b', 'B.Sc'),
        (r'\b(mba|master\s+of\s+business\s+administration)\b', 'MBA'),
        (r'\b(b\.?e\.?|bachelor\s+of\s+engineering)\b', 'B.E.'),
        (r'\b(m\.?e\.?|master\s+of\s+engineering)\b', 'M.E.'),
        (r'\b(ph\.?d\.?|doctorate|doctoral)\b', 'Ph.D.'),
        (r'\b(b\.?a\.?|bachelor\s+of\s+arts)\b', 'B.A.'),
        (r'\b(m\.?a\.?|master\s+of\s+arts)\b', 'M.A.'),
        (r'\b(b\.?com|bachelor\s+of\s+commerce)\b', 'B.Com'),
        (r'\b(m\.?com|master\s+of\s+commerce)\b', 'M.Com'),
        (r'\bhsc\b', 'HSC (Higher Secondary)'),
        (r'\bssc\b', 'SSC (Secondary School)'),
        (r'\b(diploma|associate)\b', 'Diploma'),
    ]
    
    education_found = []
    for pattern, degree_name in education_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            if degree_name not in education_found:
                education_found.append(degree_name)
    
    info["education"] = education_found
    
    # Skills extraction
    skill_keywords = [
        'python', 'javascript', 'java', 'c++', 'c#', 'rust', 'go', 'kotlin', 'swift',
        'react', 'angular', 'vue', 'node', 'express', 'django', 'flask', 'fastapi',
        'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins',
        'html', 'css', 'sql', 'git', 'linux', 'windows',
        'machine learning', 'nlp', 'computer vision', 'deep learning',
        'rest api', 'graphql', 'websocket', 'microservices', 'devops',
        'agile', 'scrum', 'jira', 'confluence', 'slack'
    ]
    for skill in skill_keywords:
        if skill in text:
            info["skills"].append(skill.title())
    info["skills"] = list(set(info["skills"]))
    
    # Clean the name if extracted
    if info["name"]:
        info["name"] = _clean_extracted_name(info["name"])
    
    return info


def _clean_extracted_name(name: str) -> str:
    """Clean extracted name - remove address, symbols, etc."""
    print(f"_clean_extracted_name: Cleaning name '{name}'")
    if not name or name == "Unknown":
        return "Unknown"
    
    # Stop at common separators
    separators = ['+', '#', '|', '@', '(', ')', ',', ':', ';']
    for sep in separators:
        if sep in name:
            name = name.split(sep)[0]
    
    # Remove numbers (likely address or phone)
    name = re.sub(r'\d+', '', name)
    
    # Remove extra whitespace and clean
    name = re.sub(r'\s+', ' ', name).strip()
    
    # Remove common non-name words
    non_name_words = ['l-', 'apt', 'street', 'road', 'nagar', 'vadodara', 'india', 'address', 'taif']
    name_lower = name.lower()
    for word in non_name_words:
        if word in name_lower:
            idx = name_lower.find(word)
            name = name[:idx].strip()
            name_lower = name.lower()
    
    # Ensure it looks like a name (mostly letters)
    if name:
        letter_count = sum(1 for c in name if c.isalpha() or c.isspace())
        if len(name) > 0 and letter_count / len(name) < 0.8:
            return "Unknown"
    
    # Title case and limit to reasonable length
    name = name.strip().title()
    words = name.split()
    if len(words) > 5:
        name = ' '.join(words[:5])
    
    return name if len(name) >= 2 else "Unknown"
